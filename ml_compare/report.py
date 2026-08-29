"""Builds Excel (.xlsx) and Word (.docx) exports from an already-computed
/api/ml_compare response payload. Takes the JSON dict as-is -- no
re-training, so exporting never re-runs cross-validation -- and reuses
pls/report.py's generic table-building helpers rather than duplicating them.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook

from i18n import DEFAULT_LANG, t
from pls.report import _add_heading, _add_table, _autofit, _fmt, _sheet_name, _write_table

COMPARISON_HEADER_KEYS = ["rpt_ml_target", "rpt_ml_predictor", "rpt_ml_sem_coef"]


def _algo_label(algo_id: str, lang: str) -> str:
    return t(f"ml_algo_{algo_id}", lang)


def _fmt_stat(stat: dict | None, digits: int = 3) -> str:
    if not stat:
        return "—"
    return f"{_fmt(stat.get('mean'), digits)} (±{_fmt(stat.get('std'), digits)})"


def _comparison_rows(data: dict, fmt_fn=lambda v: v) -> list[list]:
    rows = []
    for tr in data["targets"]:
        for p in tr["predictors"]:
            row = [tr["target_name"], p["name"], fmt_fn(p.get("sem_coefficient"))]
            for algo_id in data["algorithms"]:
                ao = tr["algorithms"].get(algo_id)
                pi = ao["permutation_importance"].get(p["id"]) if ao else None
                row.append(_fmt_stat(pi))
            rows.append(row)
    return rows


def _metrics_table(data: dict, algo_id: str, lang: str) -> tuple[list[str], list[list], bool]:
    is_classification = any(
        tr["algorithms"].get(algo_id, {}).get("task") == "classification" for tr in data["targets"]
    )
    if is_classification:
        headers = [t("rpt_ml_target", lang), t("rpt_ml_accuracy", lang), t("rpt_ml_auc", lang)]
    else:
        headers = [t("rpt_ml_target", lang), "R²", t("rpt_ml_rmse", lang)]
    rows = []
    for tr in data["targets"]:
        ao = tr["algorithms"].get(algo_id)
        if not ao:
            continue
        if is_classification:
            rows.append([tr["target_name"], _fmt_stat(ao["metrics"].get("accuracy")), _fmt_stat(ao["metrics"].get("auc"))])
        else:
            rows.append([tr["target_name"], _fmt_stat(ao["metrics"].get("r2")), _fmt_stat(ao["metrics"].get("rmse"))])
    return headers, rows, is_classification


def _importance_rows(data: dict, algo_id: str, fmt_fn=lambda v: v) -> list[list]:
    rows = []
    for tr in data["targets"]:
        ao = tr["algorithms"].get(algo_id)
        if not ao:
            continue
        native = ao.get("native_importance")
        for p in tr["predictors"]:
            native_v = native.get(p["id"]) if native else None
            perm = ao["permutation_importance"].get(p["id"])
            rows.append([tr["target_name"], p["name"], fmt_fn(native_v), _fmt_stat(perm)])
    return rows


def build_excel_report(data: dict, lang: str = DEFAULT_LANG) -> io.BytesIO:
    wb = Workbook()

    ws = wb.active
    ws.title = _sheet_name(t("rpt_sheet_overview", lang))
    overview_rows = [
        [t("rpt_method", lang), "PLS-SEM" if data["method"] == "pls" else "CB-SEM"],
        [t("rpt_ml_k", lang), data["k"]],
        [t("rpt_ml_algorithms_selected", lang), ", ".join(_algo_label(a, lang) for a in data["algorithms"])],
        [t("rpt_export_date", lang), datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]
    r = _write_table(ws, 1, [t("rpt_model_info", lang), t("rpt_value", lang)], overview_rows,
                      title=t("rpt_title_ml", lang))
    target_rows = [[tr["target_name"], tr["n_obs"], len(tr["predictors"])] for tr in data["targets"]]
    _write_table(
        ws, r, [t("rpt_ml_target", lang), t("rpt_n_obs", lang), t("rpt_ml_predictor", lang) + " (n)"],
        target_rows, title=t("rpt_ml_targets_title", lang),
    )
    _autofit(ws)

    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_ml_comparison", lang)))
    headers = [t(k, lang) for k in COMPARISON_HEADER_KEYS] + [_algo_label(a, lang) for a in data["algorithms"]]
    _write_table(ws, 1, headers, _comparison_rows(data), title=t("rpt_ml_comparison_title", lang))
    _autofit(ws)

    for algo_id in data["algorithms"]:
        ws = wb.create_sheet(_sheet_name(_algo_label(algo_id, lang)))
        m_headers, m_rows, is_cls = _metrics_table(data, algo_id, lang)
        r = _write_table(ws, 1, m_headers, m_rows, title=_algo_label(algo_id, lang))
        if is_cls:
            ws.cell(row=r, column=1, value=t("rpt_ml_logreg_note", lang))
            r += 2
        i_headers = [t("rpt_ml_target", lang), t("rpt_ml_predictor", lang), t("rpt_ml_native_importance", lang),
                     t("rpt_ml_permutation_importance", lang)]
        _write_table(ws, r, i_headers, _importance_rows(data, algo_id), title=t("rpt_ml_importance_title", lang))
        _autofit(ws)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def build_word_report(data: dict, lang: str = DEFAULT_LANG) -> io.BytesIO:
    doc = Document()
    title = doc.add_heading(t("rpt_title_ml", lang), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"{'PLS-SEM' if data['method'] == 'pls' else 'CB-SEM'} · k={data['k']} · "
        f"{t('rpt_export_date', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    _add_heading(doc, t("rpt_ml_algorithms_selected", lang), level=2)
    doc.add_paragraph(", ".join(_algo_label(a, lang) for a in data["algorithms"]))

    _add_heading(doc, t("rpt_ml_comparison_title", lang), level=1)
    headers = [t(k, lang) for k in COMPARISON_HEADER_KEYS] + [_algo_label(a, lang) for a in data["algorithms"]]
    _add_table(doc, headers, _comparison_rows(data, fmt_fn=_fmt))

    for algo_id in data["algorithms"]:
        _add_heading(doc, _algo_label(algo_id, lang), level=1)
        m_headers, m_rows, is_cls = _metrics_table(data, algo_id, lang)
        _add_table(doc, m_headers, m_rows)
        if is_cls:
            note = doc.add_paragraph()
            note_run = note.add_run(t("rpt_ml_logreg_note", lang))
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)
        i_headers = [t("rpt_ml_target", lang), t("rpt_ml_predictor", lang), t("rpt_ml_native_importance", lang),
                     t("rpt_ml_permutation_importance", lang)]
        _add_table(doc, i_headers, _importance_rows(data, algo_id, fmt_fn=_fmt))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
