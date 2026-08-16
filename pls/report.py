"""Builds Excel (.xlsx) and Word (.docx) report exports from an already-computed
/api/analyze response payload. Takes the JSON dict as-is — no re-estimation —
so exporting never re-runs bootstrapping/blindfolding. All labels are drawn
from the shared i18n catalog (see i18n.py) according to the `lang` argument,
which is independent of whatever language the original analysis used.
"""

from __future__ import annotations

import base64
import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from i18n import DEFAULT_LANG, t

DIAGRAM_EXCEL_MAX_WIDTH_PX = 900


def _decode_diagram_image(data: dict) -> bytes | None:
    """The frontend sends the results diagram as a `canvas.toDataURL()` PNG
    data URL — decode it back to raw bytes for embedding in the report, or
    None if the caller didn't include one (e.g. older cached payloads)."""
    raw = data.get("diagram_image")
    if not raw or not isinstance(raw, str):
        return None
    if "," in raw:
        raw = raw.split(",", 1)[1]
    try:
        return base64.b64decode(raw)
    except (ValueError, TypeError):
        return None

HEADER_FILL = PatternFill(start_color="EEF1FD", end_color="EEF1FD", fill_type="solid")
HEADER_FONT = Font(bold=True)


def _id_to_name(data: dict) -> dict:
    return {c["id"]: c["name"] for c in data["constructs"]}


def _fmt(v, digits: int = 3) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.{digits}f}"
    except (TypeError, ValueError):
        return str(v)


def r2_label(v, lang: str) -> str:
    if v is None:
        return "—"
    if v < 0.19:
        return t("lbl_r2_weak", lang)
    if v < 0.33:
        return t("lbl_r2_moderate", lang)
    if v < 0.67:
        return t("lbl_r2_substantial", lang)
    return t("lbl_r2_strong", lang)


def f2_label(v, lang: str) -> str:
    if v is None:
        return "—"
    if v < 0.02:
        return t("lbl_f2_none", lang)
    if v < 0.15:
        return t("lbl_f2_small", lang)
    if v < 0.35:
        return t("lbl_f2_medium", lang)
    return t("lbl_f2_large", lang)


# Moderation (interaction-term) f² uses much smaller thresholds than main-effect
# f² (Kenny 2018; Aguinis, Beaty, Boik & Pierce 2005) — interaction effects are
# inherently harder to detect, so the same 0.02/0.15/0.35 bar would mislabel
# virtually every real moderation effect as negligible.
def f2_moderation_label(v, lang: str) -> str:
    if v is None:
        return "—"
    if v < 0.005:
        return t("lbl_f2_none", lang)
    if v < 0.01:
        return t("lbl_f2_small", lang)
    if v < 0.025:
        return t("lbl_f2_medium", lang)
    return t("lbl_f2_large", lang)


def q2_label(v, lang: str) -> str:
    if v is None:
        return "—"
    if v <= 0:
        return t("lbl_q2_none", lang)
    if v < 0.02:
        return t("lbl_f2_none", lang)
    if v < 0.15:
        return t("lbl_f2_small", lang)
    if v < 0.35:
        return t("lbl_f2_medium", lang)
    return t("lbl_f2_large", lang)


def sig_label(significant, lang: str) -> str:
    if significant is None:
        return "—"
    return t("rpt_significant", lang) if significant else t("rpt_not_significant", lang)


def cmb_label(v, threshold: float, lang: str) -> str:
    if v is None:
        return "—"
    return t("lbl_cmb_warn", lang) if v > threshold else t("lbl_cmb_ok", lang)


def _total_effects_rows(data: dict, fmt_fn=lambda v: v) -> list:
    return [
        [f'{e["source_name"]} → {e["target_name"]}', fmt_fn(e["direct"]), fmt_fn(e["indirect"]), fmt_fn(e["total"])]
        for e in (data.get("structural", {}).get("total_effects") or [])
    ]


def _cmb_rows(data: dict, id_to_name: dict, lang: str, fmt_fn=lambda v: v) -> tuple[list, float]:
    cmb = data.get("common_method_bias") or {"vif": {}, "threshold": 3.3}
    threshold = cmb["threshold"]
    rows = [
        [id_to_name[cid], fmt_fn(v), cmb_label(v, threshold, lang)]
        for cid, v in cmb["vif"].items()
    ]
    return rows, threshold


def _mode_label(mode: str, lang: str) -> str:
    if mode == "I":
        return t("rpt_interaction_term", lang)
    return t("rpt_reflective", lang) if mode == "A" else t("rpt_formative", lang)


def _interaction_of_label(c: dict, id_to_name: dict) -> str:
    sources = c.get("interaction_of")
    if not sources:
        return "—"
    return " × ".join(id_to_name.get(sid, sid) for sid in sources)


def _yn(v, lang: str) -> str:
    return t("rpt_yes", lang) if v else t("rpt_no", lang)


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def _write_table(ws, start_row: int, headers: list[str], rows: list[list], title: str | None = None):
    r = start_row
    if title:
        ws.cell(row=r, column=1, value=title).font = Font(bold=True, size=12)
        r += 1
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=r, column=c, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
    r += 1
    for row in rows:
        for c, val in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=val)
        r += 1
    return r + 1  # next free row, with one blank line after the table


def _sheet_name(name: str) -> str:
    """Excel sheet names have a hard 31-character limit."""
    return name[:31]


def _autofit(ws, max_width: int = 42):
    widths: dict[int, int] = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            widths[cell.column] = min(max_width, max(widths.get(cell.column, 8), len(str(cell.value)) + 2))
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width


def build_excel_report(data: dict, lang: str = DEFAULT_LANG) -> io.BytesIO:
    id_to_name = _id_to_name(data)
    wb = Workbook()

    # --- Summary ---
    ws = wb.active
    ws.title = _sheet_name(t("rpt_sheet_overview", lang))
    rows = [
        [t("rpt_n_obs", lang), data.get("n_obs")],
        [t("rpt_converged", lang), _yn(data.get("converged"), lang)],
        [t("rpt_n_iterations", lang), data.get("iterations")],
    ]
    boot = data.get("bootstrap")
    if boot:
        rows.append([t("rpt_bootstrap_requested", lang), boot.get("requested")])
        rows.append([t("rpt_bootstrap_valid", lang), boot.get("valid")])
    rows.append([t("rpt_export_date", lang), datetime.now().strftime("%Y-%m-%d %H:%M")])
    r = _write_table(ws, 1, [t("rpt_model_info", lang), t("rpt_value", lang)], rows,
                      title=t("rpt_title_pls", lang))

    construct_rows = [
        [c["name"], _mode_label(c["mode"], lang), ", ".join(c["indicators"]), _yn(c["is_endogenous"], lang),
         _interaction_of_label(c, id_to_name)]
        for c in data["constructs"]
    ]
    r = _write_table(ws, r, [t("rpt_construct", lang), t("rpt_measurement_type", lang),
                              t("rpt_indicators", lang), t("rpt_endogenous", lang), t("rpt_moderation_of", lang)],
                      construct_rows, title=t("rpt_construct_list", lang))
    _autofit(ws)

    # --- Path diagram (rendered client-side, sent as a PNG data URL) ---
    diagram_bytes = _decode_diagram_image(data)
    if diagram_bytes:
        ws = wb.create_sheet(_sheet_name(t("rpt_sheet_diagram", lang)))
        xl_img = XLImage(io.BytesIO(diagram_bytes))
        if xl_img.width > DIAGRAM_EXCEL_MAX_WIDTH_PX:
            scale = DIAGRAM_EXCEL_MAX_WIDTH_PX / xl_img.width
            xl_img.width = int(xl_img.width * scale)
            xl_img.height = int(xl_img.height * scale)
        ws.add_image(xl_img, "A1")

    # --- Outer loadings ---
    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_outer_loadings", lang)))
    m = data["measurement"]
    boot_load = {row["indicator"]: row for row in (m.get("outer_loadings_bootstrap") or [])}
    headers = [t("rpt_indicator", lang), t("rpt_construct", lang), t("rpt_outer_loading", lang),
               t("rpt_outer_weight", lang)]
    has_boot = bool(m.get("outer_loadings_bootstrap"))
    if has_boot:
        headers += [t("rpt_stdev", lang), t("rpt_t_stat", lang), t("rpt_p_value", lang), t("rpt_significance", lang)]
    rows = []
    for c in data["constructs"]:
        for ind in c["indicators"]:
            row = [ind, c["name"], m["outer_loadings"].get(ind), m["outer_weights"].get(ind)]
            if has_boot:
                b = boot_load.get(ind)
                row += [b["std"], b["t_stat"], b["p_value"], sig_label(b["significant"], lang)] if b else ["—"] * 4
            rows.append(row)
    _write_table(ws, 1, headers, rows, title=f"{t('rpt_outer_loading', lang)} & {t('rpt_outer_weight', lang)}")
    _autofit(ws)

    # --- Cross loadings ---
    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_cross_loadings", lang)))
    cl = m["cross_loadings"]
    # Interaction/moderation constructs have no indicators of their own, so they
    # never appear as a column in the cross-loadings matrix (see htmt()'s docstring
    # for the same reasoning) — derive the column set from `cl` itself rather than
    # from the full construct list, which would include them.
    cl_ids = list(cl.keys())
    rows = []
    for c in data["constructs"]:
        for ind in c["indicators"]:
            rows.append([ind, c["name"]] + [cl[cid][ind] for cid in cl_ids])
    _write_table(ws, 1, [t("rpt_indicator", lang), t("rpt_construct", lang)] + [id_to_name[cid] for cid in cl_ids],
                 rows, title=t("rpt_sheet_cross_loadings", lang))
    _autofit(ws)

    # --- Reliability & validity ---
    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_reliability", lang)))
    rows = []
    for c in data["constructs"]:
        if c["mode"] == "I":
            rows.append([c["name"], "—", "—", "—", "—", t("rpt_moderation_no_reliability", lang)])
            continue
        if c["mode"] != "A":
            rows.append([c["name"], "—", "—", "—", "—", t("rpt_formative_no_reliability", lang)])
            continue
        cid = c["id"]
        rows.append([
            c["name"], m["cronbachs_alpha"].get(cid), m["rho_a"].get(cid),
            m["composite_reliability"].get(cid), m["ave"].get(cid), "",
        ])
    _write_table(ws, 1, [t("rpt_construct", lang), t("rpt_cronbachs_alpha", lang), t("rpt_rho_a", lang),
                          t("rpt_composite_reliability", lang), t("rpt_ave", lang), t("rpt_note", lang)],
                 rows, title=t("rpt_sheet_reliability", lang))
    _autofit(ws)

    # --- Discriminant validity ---
    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_discriminant", lang)))
    dv = data["discriminant_validity"]

    def matrix_rows(matrix, ids):
        return [[id_to_name[rid]] + [matrix[cid][rid] for cid in ids] for rid in ids]

    # Fornell-Larcker only covers reflective (Mode A) constructs — AVE (its
    # diagonal) isn't defined for formative ones — so its id set can be a
    # strict subset of all constructs; HTMT has no such restriction.
    fl_ids = list(dv["fornell_larcker"].keys())
    htmt_ids = list(dv["htmt"].keys())
    r = _write_table(ws, 1, [""] + [id_to_name[cid] for cid in fl_ids],
                      matrix_rows(dv["fornell_larcker"], fl_ids), title=t("rpt_fornell_larcker", lang))
    _write_table(ws, r, [""] + [id_to_name[cid] for cid in htmt_ids],
                 matrix_rows(dv["htmt"], htmt_ids), title=t("rpt_htmt", lang))
    _autofit(ws)

    # --- Structural model ---
    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_structural", lang)))
    st = data["structural"]
    has_boot_path = any("t_stat" in p for p in st["paths"])
    headers = [t("rpt_path", lang), t("rpt_path_coefficient", lang)]
    if has_boot_path:
        headers += [t("rpt_stdev", lang), t("rpt_t_stat", lang), t("rpt_p_value", lang), t("rpt_significance", lang)]
    headers += [t("rpt_f_squared", lang), t("rpt_f2_effect", lang)]
    rows = []
    for p in st["paths"]:
        label_fn = f2_moderation_label if p.get("is_interaction") else f2_label
        row = [f'{p["source_name"]} → {p["target_name"]}', p["coefficient"]]
        if has_boot_path:
            row += [p.get("bootstrap_std"), p.get("t_stat"), p.get("p_value"), sig_label(p.get("significant"), lang)]
        row += [p["f_squared"], label_fn(p["f_squared"], lang)]
        rows.append(row)
    r = _write_table(ws, 1, headers, rows, title=t("rpt_path_coefficients", lang))

    te_rows = _total_effects_rows(data)
    if te_rows:
        r = _write_table(
            ws, r, [t("rpt_path", lang), t("rpt_direct_effect", lang), t("rpt_indirect_effect", lang),
                    t("rpt_total_effect", lang)],
            te_rows, title=t("rpt_total_effects_title", lang),
        )
        ws.cell(row=r, column=1, value=t("rpt_total_effects_note", lang))
        r += 2

    r2_rows = []
    for cid, r2 in st["r_squared"].items():
        q2 = st.get("q_squared", {}).get(cid)
        skip_reason = st.get("q_squared_skipped", {}).get(cid)
        r2_rows.append([
            id_to_name[cid], r2, st["r_squared_adj"].get(cid), r2_label(r2, lang),
            q2 if q2 is not None else (skip_reason or "—"),
            q2_label(q2, lang) if q2 is not None else "—",
        ])
    r = _write_table(
        ws, r,
        [t("rpt_endogenous_construct", lang), t("rpt_r2", lang), t("rpt_r2_adj", lang), t("rpt_r2_assessment", lang),
         t("rpt_q2", lang, d=st.get("omission_distance", "—")), t("rpt_q2_assessment", lang)],
        r2_rows, title=t("rpt_r2_q2_title", lang),
    )

    vif_rows = []
    inner = st["inner_vif"]
    for target_id, preds in inner.items():
        for pred_id, v in preds.items():
            if v is not None:
                vif_rows.append([
                    f"{id_to_name[pred_id]} → {id_to_name[target_id]}{t('rpt_structural_suffix', lang)}", v,
                ])
    for ind, v in (m.get("outer_vif") or {}).items():
        vif_rows.append([f"{ind}{t('rpt_formative_measurement_suffix', lang)}", v])
    r = _write_table(ws, r, [t("rpt_pair", lang), t("rpt_vif", lang)], vif_rows, title=t("rpt_vif_title", lang))

    cmb_rows, cmb_threshold = _cmb_rows(data, id_to_name, lang)
    r = _write_table(ws, r, [t("rpt_construct", lang), t("rpt_vif", lang), t("rpt_cmb_assessment", lang)],
                      cmb_rows, title=t("rpt_cmb_title", lang))
    ws.cell(row=r, column=1, value=t("rpt_cmb_note", lang, threshold=cmb_threshold))
    _autofit(ws)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x34, 0x57, 0xD5)
    return h


def _add_table(doc: Document, headers: list[str], rows: list[list], col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "—" if val is None else str(val)
    doc.add_paragraph()
    return table


def build_word_report(data: dict, lang: str = DEFAULT_LANG) -> io.BytesIO:
    id_to_name = _id_to_name(data)
    m = data["measurement"]
    dv = data["discriminant_validity"]
    st = data["structural"]

    doc = Document()
    title = doc.add_heading(t("rpt_title_pls", lang), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    conv_text = t("rpt_word_converged", lang) if data.get("converged") else t("rpt_word_not_converged", lang)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"n = {data.get('n_obs')} · {conv_text} {t('rpt_after_iterations', lang, it=data.get('iterations'))} · "
        f"{t('rpt_export_date', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    # --- Model overview ---
    _add_heading(doc, t("rpt_section_overview", lang), level=1)
    _add_table(
        doc, [t("rpt_construct", lang), t("rpt_measurement_type", lang), t("rpt_indicators", lang),
              t("rpt_endogenous", lang), t("rpt_moderation_of", lang)],
        [[c["name"], _mode_label(c["mode"], lang), len(c["indicators"]), _yn(c["is_endogenous"], lang),
          _interaction_of_label(c, id_to_name)]
         for c in data["constructs"]],
    )

    diagram_bytes = _decode_diagram_image(data)
    if diagram_bytes:
        _add_heading(doc, t("rpt_section_diagram", lang), level=1)
        doc.add_picture(io.BytesIO(diagram_bytes), width=Inches(6.3))

    # --- Measurement model ---
    _add_heading(doc, t("rpt_section_measurement", lang), level=1)
    _add_heading(doc, t("rpt_section_loadings", lang), level=2)
    boot_load = {row["indicator"]: row for row in (m.get("outer_loadings_bootstrap") or [])}
    has_boot = bool(m.get("outer_loadings_bootstrap"))
    headers = [t("rpt_indicator", lang), t("rpt_construct", lang), t("rpt_outer_loading", lang),
               t("rpt_outer_weight", lang)]
    if has_boot:
        headers += [t("rpt_stdev", lang), t("rpt_z_short", lang).upper(), t("rpt_p_short", lang).upper(),
                    t("rpt_significance", lang)]
    rows = []
    for c in data["constructs"]:
        for ind in c["indicators"]:
            row = [ind, c["name"], _fmt(m["outer_loadings"].get(ind)), _fmt(m["outer_weights"].get(ind))]
            if has_boot:
                b = boot_load.get(ind)
                row += [_fmt(b["std"]), _fmt(b["t_stat"], 2), _fmt(b["p_value"], 4),
                        sig_label(b["significant"], lang)] if b else ["—"] * 4
            rows.append(row)
    _add_table(doc, headers, rows)

    _add_heading(doc, t("rpt_section_reliability", lang), level=2)
    rel_rows = []
    for c in data["constructs"]:
        if c["mode"] != "A":
            rel_rows.append([c["name"], "—", "—", "—", "—"])
            continue
        cid = c["id"]
        rel_rows.append([
            c["name"], _fmt(m["cronbachs_alpha"].get(cid)), _fmt(m["rho_a"].get(cid)),
            _fmt(m["composite_reliability"].get(cid)), _fmt(m["ave"].get(cid)),
        ])
    _add_table(doc, [t("rpt_construct", lang), t("rpt_cronbachs_alpha", lang), t("rpt_rho_a", lang),
                      t("rpt_composite_reliability", lang), t("rpt_ave", lang)], rel_rows)

    # Fornell-Larcker only covers reflective (Mode A) constructs (AVE, its
    # diagonal, isn't defined for formative ones), so its id set can be a
    # strict subset of all constructs; HTMT has no such restriction.
    fl_ids = list(dv["fornell_larcker"].keys())
    htmt_ids = list(dv["htmt"].keys())
    _add_heading(doc, t("rpt_section_fl", lang), level=2)
    _add_table(
        doc, [""] + [id_to_name[cid] for cid in fl_ids],
        [[id_to_name[rid]] + [_fmt(dv["fornell_larcker"][cid][rid]) for cid in fl_ids] for rid in fl_ids],
    )
    _add_heading(doc, t("rpt_section_htmt", lang), level=2)
    _add_table(
        doc, [""] + [id_to_name[cid] for cid in htmt_ids],
        [[id_to_name[rid]] + [_fmt(dv["htmt"][cid][rid]) for cid in htmt_ids] for rid in htmt_ids],
    )

    # --- Structural model ---
    _add_heading(doc, t("rpt_section_structural", lang), level=1)
    _add_heading(doc, t("rpt_section_path_coef", lang), level=2)
    has_boot_path = any("t_stat" in p for p in st["paths"])
    headers = [t("rpt_path", lang), "β"]
    if has_boot_path:
        headers += [t("rpt_stdev", lang), t("rpt_z_short", lang).upper(), t("rpt_p_short", lang).upper(),
                    t("rpt_significance", lang)]
    headers += [t("rpt_f_squared", lang), t("rpt_f2_effect", lang)]
    rows = []
    for p in st["paths"]:
        label_fn = f2_moderation_label if p.get("is_interaction") else f2_label
        row = [f'{p["source_name"]} → {p["target_name"]}', _fmt(p["coefficient"])]
        if has_boot_path:
            row += [_fmt(p.get("bootstrap_std")), _fmt(p.get("t_stat"), 2), _fmt(p.get("p_value"), 4),
                    sig_label(p.get("significant"), lang)]
        row += [_fmt(p["f_squared"]), label_fn(p["f_squared"], lang)]
        rows.append(row)
    _add_table(doc, headers, rows)

    te_rows = _total_effects_rows(data, fmt_fn=_fmt)
    if te_rows:
        _add_heading(doc, t("rpt_total_effects_title", lang), level=2)
        _add_table(doc, [t("rpt_path", lang), t("rpt_direct_effect", lang), t("rpt_indirect_effect", lang),
                          t("rpt_total_effect", lang)], te_rows)
        te_note = doc.add_paragraph()
        te_note_run = te_note.add_run(t("rpt_total_effects_note", lang))
        te_note_run.font.size = Pt(9)
        te_note_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    _add_heading(doc, t("rpt_section_r2q2", lang), level=2)
    r2_rows = []
    for cid, r2 in st["r_squared"].items():
        q2 = st.get("q_squared", {}).get(cid)
        skip_reason = st.get("q_squared_skipped", {}).get(cid)
        r2_rows.append([
            id_to_name[cid], _fmt(r2), _fmt(st["r_squared_adj"].get(cid)), r2_label(r2, lang),
            _fmt(q2) if q2 is not None else (skip_reason or "—"),
            q2_label(q2, lang) if q2 is not None else "—",
        ])
    _add_table(doc, [t("rpt_construct", lang), t("rpt_r2", lang), t("rpt_r2_adj", lang),
                      t("rpt_r2_assessment", lang), "Q²", t("rpt_q2_assessment", lang)], r2_rows)

    _add_heading(doc, t("rpt_section_vif", lang), level=2)
    vif_rows = []
    for target_id, preds in st["inner_vif"].items():
        for pred_id, v in preds.items():
            if v is not None:
                vif_rows.append([
                    f"{id_to_name[pred_id]} → {id_to_name[target_id]}{t('rpt_structural_suffix', lang)}", _fmt(v),
                ])
    for ind, v in (m.get("outer_vif") or {}).items():
        vif_rows.append([f"{ind}{t('rpt_formative_measurement_suffix', lang)}", _fmt(v)])
    if vif_rows:
        _add_table(doc, [t("rpt_pair", lang), t("rpt_vif", lang)], vif_rows)
    else:
        doc.add_paragraph(t("rpt_no_vif_pairs", lang))

    _add_heading(doc, t("rpt_cmb_title", lang), level=2)
    cmb_rows, cmb_threshold = _cmb_rows(data, id_to_name, lang, fmt_fn=_fmt)
    _add_table(doc, [t("rpt_construct", lang), t("rpt_vif", lang), t("rpt_cmb_assessment", lang)], cmb_rows)
    cmb_note = doc.add_paragraph()
    cmb_note_run = cmb_note.add_run(t("rpt_cmb_note", lang, threshold=cmb_threshold))
    cmb_note_run.font.size = Pt(9)
    cmb_note_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    if data.get("bootstrap"):
        note = doc.add_paragraph()
        run = note.add_run(t(
            "rpt_bootstrap_note", lang,
            requested=data["bootstrap"]["requested"], valid=data["bootstrap"]["valid"],
        ))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)
    if st.get("omission_distance"):
        note2 = doc.add_paragraph()
        run2 = note2.add_run(t("rpt_blindfolding_note", lang, d=st["omission_distance"]))
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
