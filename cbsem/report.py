"""Builds Excel (.xlsx) and Word (.docx) report exports from an already-computed
/api/analyze_cbsem response payload. Mirrors pls/report.py's structure, adapted
for CB-SEM's own reporting conventions (fit indices instead of blindfolding,
ML standard errors/z/p instead of bootstrap SE/t/p). Labels come from the
shared i18n catalog according to the `lang` argument.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from i18n import DEFAULT_LANG, t
from pls.report import (
    DIAGRAM_EXCEL_MAX_WIDTH_PX,
    _cmb_rows,
    _decode_diagram_image,
    _interaction_of_label,
    _specific_indirect_rows,
    _total_effects_rows,
)

HEADER_FILL = PatternFill(start_color="EEF1FD", end_color="EEF1FD", fill_type="solid")
HEADER_FONT = Font(bold=True)

FIT_INDEX_KEYS = [
    "rpt_fit_chi_square", "rpt_fit_df", "rpt_fit_chi2_p", "rpt_fit_cfi", "rpt_fit_tli",
    "rpt_fit_rmsea", "rpt_fit_srmr", "rpt_fit_gfi", "rpt_fit_agfi", "rpt_fit_nfi",
    "rpt_fit_aic", "rpt_fit_bic",
]
FIT_INDEX_DATA_KEYS = {
    "rpt_fit_chi_square": "chi_square", "rpt_fit_df": "df", "rpt_fit_chi2_p": "chi_square_p_value",
    "rpt_fit_cfi": "cfi", "rpt_fit_tli": "tli", "rpt_fit_rmsea": "rmsea", "rpt_fit_srmr": "srmr",
    "rpt_fit_gfi": "gfi", "rpt_fit_agfi": "agfi", "rpt_fit_nfi": "nfi", "rpt_fit_aic": "aic",
    "rpt_fit_bic": "bic",
}


def _id_to_name(data: dict) -> dict:
    return {c["id"]: c["name"] for c in data["constructs"]}


def sig_label(p_value, lang: str) -> str:
    if p_value is None:
        return "—"
    return t("rpt_significant", lang) if p_value < 0.05 else t("rpt_not_significant", lang)


def fit_verdict(data_key: str, v, lang: str) -> str:
    if v is None:
        return ""
    thresholds = {
        "cfi": (v >= 0.90, v >= 0.95),
        "tli": (v >= 0.90, v >= 0.95),
        "rmsea": (v <= 0.08, v <= 0.06),
        "srmr": (v <= 0.08, v <= 0.05),
        "gfi": (v >= 0.90, v >= 0.95),
        "nfi": (v >= 0.90, v >= 0.95),
    }
    if data_key not in thresholds:
        return ""
    acceptable, good = thresholds[data_key]
    if good:
        return t("lbl_fit_good", lang)
    if acceptable:
        return t("lbl_fit_acceptable", lang)
    return t("lbl_fit_poor", lang)


def _mode_label(mode: str, lang: str, calc_method: str | None = None) -> str:
    if mode == "I":
        base = t("rpt_interaction_term", lang)
        # CB-SEM ignores calc_method and always treats interactions as two-stage
        # (see cbsem/moderation.py) regardless of what was configured for them.
        return f"{base} ({t('rpt_calc_method_two_stage', lang)})"
    return t("rpt_reflective", lang)  # CB-SEM only supports reflective measurement otherwise


def _yn(v, lang: str) -> str:
    return t("rpt_yes", lang) if v else t("rpt_no", lang)


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def _write_table(ws, start_row, headers, rows, title=None):
    r = start_row
    if title:
        ws.cell(row=r, column=1, value=title).font = Font(bold=True, size=12)
        r += 1
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=r, column=c, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
    r += 1
    for row in rows:
        for c, val in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=val)
        r += 1
    return r + 1


def _sheet_name(name: str) -> str:
    """Excel sheet names have a hard 31-character limit."""
    return name[:31]


def _autofit(ws, max_width=42):
    widths = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            widths[cell.column] = min(max_width, max(widths.get(cell.column, 8), len(str(cell.value)) + 2))
    for col, width in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = width


def build_excel_report(data: dict, lang: str = DEFAULT_LANG) -> io.BytesIO:
    id_to_name = _id_to_name(data)
    wb = Workbook()

    ws = wb.active
    ws.title = _sheet_name(t("rpt_sheet_overview", lang))
    rows = [
        [t("rpt_method", lang), "CB-SEM (Maximum Likelihood)"],
        [t("rpt_n_obs", lang), data.get("n_obs")],
        [t("rpt_converged", lang), _yn(data.get("converged"), lang)],
        [t("rpt_optimizer_message", lang), data.get("optimizer_message")],
        [t("rpt_export_date", lang), datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]
    r = _write_table(ws, 1, [t("rpt_model_info", lang), t("rpt_value", lang)], rows,
                      title=t("rpt_title_cbsem", lang))

    construct_rows = [
        [c["name"], _mode_label(c["mode"], lang, c.get("calc_method")), ", ".join(c["indicators"]), _yn(c["is_endogenous"], lang),
         _interaction_of_label(c, id_to_name)]
        for c in data["constructs"]
    ]
    r = _write_table(ws, r, [t("rpt_construct", lang), t("rpt_measurement_type", lang),
                              t("rpt_indicators", lang), t("rpt_endogenous", lang), t("rpt_moderation_of", lang)],
                      construct_rows, title=t("rpt_construct_list", lang))

    fit_rows = []
    fi = data["fit_indices"]
    for key in FIT_INDEX_KEYS:
        data_key = FIT_INDEX_DATA_KEYS[key]
        fit_rows.append([t(key, lang), fi.get(data_key), fit_verdict(data_key, fi.get(data_key), lang)])
    _write_table(ws, r, [t("rpt_fit_index", lang), t("rpt_value", lang), t("rpt_fit_assessment", lang)],
                 fit_rows, title=t("rpt_model_fit", lang))
    _autofit(ws)

    diagram_bytes = _decode_diagram_image(data)
    if diagram_bytes:
        ws = wb.create_sheet(_sheet_name(t("rpt_sheet_diagram", lang)))
        xl_img = XLImage(io.BytesIO(diagram_bytes))
        if xl_img.width > DIAGRAM_EXCEL_MAX_WIDTH_PX:
            scale = DIAGRAM_EXCEL_MAX_WIDTH_PX / xl_img.width
            xl_img.width = int(xl_img.width * scale)
            xl_img.height = int(xl_img.height * scale)
        ws.add_image(xl_img, "A1")

    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_measurement", lang)))
    m = data["measurement"]
    headers = [t("rpt_indicator", lang), t("rpt_construct", lang), t("rpt_unstandardized", lang),
               t("rpt_standardized", lang), t("rpt_se", lang), t("rpt_z_value", lang), t("rpt_p_value", lang),
               t("rpt_note", lang)]
    rows = []
    for row in m["loadings"]:
        note = t("rpt_reference_indicator", lang) if row["is_reference"] else sig_label(row["p"], lang)
        rows.append([row["indicator"], id_to_name[row["construct"]], row["unstd"], row["std"],
                     row["se"], row["z"], row["p"], note])
    r = _write_table(ws, 1, headers, rows, title=t("rpt_factor_loadings", lang))

    rel_rows = []
    for c in data["constructs"]:
        cid = c["id"]
        rel_rows.append([
            c["name"], m["cronbachs_alpha"].get(cid), m["composite_reliability"].get(cid), m["ave"].get(cid),
        ])
    _write_table(ws, r, [t("rpt_construct", lang), t("rpt_cronbachs_alpha", lang),
                          t("rpt_composite_reliability", lang), t("rpt_ave", lang)],
                 rel_rows, title=t("rpt_sheet_reliability", lang))
    _autofit(ws)

    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_discriminant", lang)))
    dv = data["discriminant_validity"]
    fl_ids = list(dv["fornell_larcker"].keys())
    htmt_ids = list(dv["htmt"].keys())

    def matrix_rows(matrix, ids):
        return [[id_to_name[rid]] + [matrix[cid][rid] for cid in ids] for rid in ids]

    r = _write_table(ws, 1, [""] + [id_to_name[cid] for cid in fl_ids],
                      matrix_rows(dv["fornell_larcker"], fl_ids), title=t("rpt_fornell_larcker", lang))
    _write_table(ws, r, [""] + [id_to_name[cid] for cid in htmt_ids],
                 matrix_rows(dv["htmt"], htmt_ids), title=t("rpt_htmt", lang))
    _autofit(ws)

    ws = wb.create_sheet(_sheet_name(t("rpt_sheet_structural", lang)))
    st = data["structural"]
    headers = [t("rpt_path", lang), f"{t('rpt_unstandardized', lang)} (B)", f"{t('rpt_standardized', lang)} (β)",
               t("rpt_se", lang), t("rpt_z_value", lang), t("rpt_p_value", lang), t("rpt_significance", lang)]
    rows = []
    for p in st["paths"]:
        rows.append([
            f'{p["source_name"]} → {p["target_name"]}', p["unstd"], p["std"], p["se"], p["z"], p["p"],
            sig_label(p["p"], lang),
        ])
    r = _write_table(ws, 1, headers, rows, title=t("rpt_path_coefficients", lang))

    te_rows = _total_effects_rows(data)
    if te_rows:
        r = _write_table(
            ws, r, [t("rpt_path", lang), t("rpt_direct_effect", lang), t("rpt_indirect_effect", lang),
                    t("rpt_total_effect", lang)],
            te_rows, title=t("rpt_total_effects_title", lang),
        )
        ws.cell(row=r, column=1, value=t("rpt_total_effects_note", lang))
        r += 2

    si_rows, si_has_boot = _specific_indirect_rows(data, lang)
    if si_rows:
        si_headers = [t("rpt_path", lang), t("rpt_indirect_effect", lang)]
        if si_has_boot:
            si_headers += [t("rpt_stdev", lang), t("rpt_t_stat", lang), t("rpt_p_value", lang), t("rpt_significance", lang)]
        r = _write_table(ws, r, si_headers, si_rows, title=t("rpt_specific_indirect_title", lang))
        ws.cell(row=r, column=1, value=t("rpt_specific_indirect_note", lang))
        r += 2

    r2_rows = [[id_to_name[cid], r2] for cid, r2 in st["r_squared"].items()]
    r = _write_table(ws, r, [t("rpt_endogenous_construct", lang), t("rpt_r2", lang)], r2_rows,
                      title=t("rpt_r2_only_title", lang))

    cmb_rows, cmb_threshold = _cmb_rows(data, id_to_name, lang)
    r = _write_table(ws, r, [t("rpt_construct", lang), t("rpt_vif", lang), t("rpt_cmb_assessment", lang)],
                      cmb_rows, title=t("rpt_cmb_title", lang))
    ws.cell(row=r, column=1, value=t("rpt_cmb_note", lang, threshold=cmb_threshold))
    _autofit(ws)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def _fmt(v, digits=3):
    if v is None:
        return "—"
    try:
        return f"{float(v):.{digits}f}"
    except (TypeError, ValueError):
        return str(v)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x34, 0x57, 0xD5)
    return h


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "—" if val is None else str(val)
    doc.add_paragraph()
    return table


def build_word_report(data: dict, lang: str = DEFAULT_LANG) -> io.BytesIO:
    id_to_name = _id_to_name(data)
    m = data["measurement"]
    dv = data["discriminant_validity"]
    st = data["structural"]
    fi = data["fit_indices"]

    doc = Document()
    title = doc.add_heading(t("rpt_title_cbsem", lang), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    conv_text = t("rpt_word_converged", lang) if data.get("converged") else t("rpt_word_not_converged", lang)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"n = {data.get('n_obs')} · {conv_text} · {t('rpt_export_date', lang)}: "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    _add_heading(doc, t("rpt_cbsem_section_overview", lang), level=1)
    _add_table(
        doc, [t("rpt_construct", lang), t("rpt_measurement_type", lang), t("rpt_indicators", lang),
              t("rpt_endogenous", lang), t("rpt_moderation_of", lang)],
        [[c["name"], _mode_label(c["mode"], lang, c.get("calc_method")), len(c["indicators"]), _yn(c["is_endogenous"], lang),
          _interaction_of_label(c, id_to_name)]
         for c in data["constructs"]],
    )

    diagram_bytes = _decode_diagram_image(data)
    if diagram_bytes:
        _add_heading(doc, t("rpt_section_diagram", lang), level=1)
        doc.add_picture(io.BytesIO(diagram_bytes), width=Inches(6.3))

    _add_heading(doc, t("rpt_cbsem_section_fit", lang), level=1)
    _add_table(
        doc, [t("rpt_fit_index", lang), t("rpt_value", lang), t("rpt_fit_assessment", lang)],
        [[t(key, lang), _fmt(fi.get(FIT_INDEX_DATA_KEYS[key]), 4 if key == "rpt_fit_chi2_p" else 3),
          fit_verdict(FIT_INDEX_DATA_KEYS[key], fi.get(FIT_INDEX_DATA_KEYS[key]), lang)]
         for key in FIT_INDEX_KEYS],
    )

    _add_heading(doc, t("rpt_cbsem_section_measurement", lang), level=1)
    _add_heading(doc, t("rpt_cbsem_section_loadings", lang), level=2)
    rows = []
    for row in m["loadings"]:
        note = t("rpt_reference_short", lang) if row["is_reference"] else sig_label(row["p"], lang)
        rows.append([row["indicator"], id_to_name[row["construct"]], _fmt(row["unstd"]), _fmt(row["std"]),
                     _fmt(row["se"]), _fmt(row["z"], 2), _fmt(row["p"], 4), note])
    _add_table(doc, [t("rpt_indicator", lang), t("rpt_construct", lang), "B", "β", t("rpt_se", lang),
                      t("rpt_z_short", lang), t("rpt_p_short", lang), t("rpt_note", lang)], rows)

    _add_heading(doc, t("rpt_cbsem_section_reliability", lang), level=2)
    rel_rows = [
        [c["name"], _fmt(m["cronbachs_alpha"].get(c["id"])), _fmt(m["composite_reliability"].get(c["id"])),
         _fmt(m["ave"].get(c["id"]))]
        for c in data["constructs"]
    ]
    _add_table(doc, [t("rpt_construct", lang), t("rpt_cronbachs_alpha", lang),
                      t("rpt_composite_reliability", lang), t("rpt_ave", lang)], rel_rows)

    fl_ids = list(dv["fornell_larcker"].keys())
    htmt_ids = list(dv["htmt"].keys())
    _add_heading(doc, t("rpt_cbsem_section_fl", lang), level=2)
    _add_table(
        doc, [""] + [id_to_name[cid] for cid in fl_ids],
        [[id_to_name[rid]] + [_fmt(dv["fornell_larcker"][cid][rid]) for cid in fl_ids] for rid in fl_ids],
    )
    _add_heading(doc, t("rpt_cbsem_section_htmt", lang), level=2)
    _add_table(
        doc, [""] + [id_to_name[cid] for cid in htmt_ids],
        [[id_to_name[rid]] + [_fmt(dv["htmt"][cid][rid]) for cid in htmt_ids] for rid in htmt_ids],
    )

    _add_heading(doc, t("rpt_cbsem_section_structural", lang), level=1)
    rows = []
    for p in st["paths"]:
        rows.append([f'{p["source_name"]} → {p["target_name"]}', _fmt(p["unstd"]), _fmt(p["std"]),
                     _fmt(p["se"]), _fmt(p["z"], 2), _fmt(p["p"], 4), sig_label(p["p"], lang)])
    _add_table(doc, [t("rpt_path", lang), "B", "β", t("rpt_se", lang), t("rpt_z_short", lang),
                      t("rpt_p_short", lang), t("rpt_significance", lang)], rows)

    te_rows = _total_effects_rows(data, fmt_fn=_fmt)
    if te_rows:
        _add_heading(doc, t("rpt_total_effects_title", lang), level=2)
        _add_table(doc, [t("rpt_path", lang), t("rpt_direct_effect", lang), t("rpt_indirect_effect", lang),
                          t("rpt_total_effect", lang)], te_rows)
        te_note = doc.add_paragraph()
        te_note_run = te_note.add_run(t("rpt_total_effects_note", lang))
        te_note_run.font.size = Pt(9)
        te_note_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    si_rows, si_has_boot = _specific_indirect_rows(data, lang, fmt_fn=_fmt)
    if si_rows:
        _add_heading(doc, t("rpt_specific_indirect_title", lang), level=2)
        si_headers = [t("rpt_path", lang), t("rpt_indirect_effect", lang)]
        if si_has_boot:
            si_headers += [t("rpt_stdev", lang), t("rpt_t_stat", lang), t("rpt_p_value", lang), t("rpt_significance", lang)]
        _add_table(doc, si_headers, si_rows)
        si_note = doc.add_paragraph()
        si_note_run = si_note.add_run(t("rpt_specific_indirect_note", lang))
        si_note_run.font.size = Pt(9)
        si_note_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    _add_heading(doc, t("rpt_cbsem_section_r2", lang), level=2)
    _add_table(doc, [t("rpt_construct", lang), t("rpt_r2", lang)],
               [[id_to_name[cid], _fmt(r2)] for cid, r2 in st["r_squared"].items()])

    _add_heading(doc, t("rpt_cmb_title", lang), level=2)
    cmb_rows, cmb_threshold = _cmb_rows(data, id_to_name, lang, fmt_fn=_fmt)
    _add_table(doc, [t("rpt_construct", lang), t("rpt_vif", lang), t("rpt_cmb_assessment", lang)], cmb_rows)
    cmb_note = doc.add_paragraph()
    cmb_note_run = cmb_note.add_run(t("rpt_cmb_note", lang, threshold=cmb_threshold))
    cmb_note_run.font.size = Pt(9)
    cmb_note_run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    note = doc.add_paragraph()
    run = note.add_run(t("rpt_r2_note", lang))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x73, 0x85)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
